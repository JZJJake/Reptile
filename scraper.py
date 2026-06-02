import os
import re
import urllib.parse
from datetime import datetime, timezone, timedelta
import asyncio
import random
import pandas as pd
from playwright.async_api import async_playwright
from playwright_stealth import Stealth
from bs4 import BeautifulSoup
import html2text
import aiohttp
from readability import Document
import io
import hashlib
import time

import db_manager

# ---------------------------------------------------------------------------
# Global state
# ---------------------------------------------------------------------------

task_events = {}          # task_id -> {'pause': Event, 'stop': Event}
task_root_prefixes = {}   # task_id -> dynamic root prefix

# Per-domain circuit breaker: domain -> {'fail_count': int, 'open_until': float}
_domain_circuit: dict = {}
CIRCUIT_THRESHOLD = 5
CIRCUIT_COOLDOWN = 300  # seconds

# Per-domain rate limiting: domain -> last request timestamp (epoch float)
_domain_last_request: dict = {}
DOMAIN_MIN_INTERVAL = 1.0  # seconds

def init_task_events(task_id: str):
    task_events[task_id] = {
        'pause': asyncio.Event(),
        'stop': asyncio.Event()
    }
    task_events[task_id]['pause'].set()
    task_events[task_id]['stop'].clear()

# ---------------------------------------------------------------------------
# Utility helpers
# ---------------------------------------------------------------------------

def sanitize_filename(name):
    if not name:
        return "Untitled_" + datetime.now().strftime("%H%M%S")
    safe_name = re.sub(r'[\\/*?:"<>|]', "", str(name))
    safe_name = re.sub(r'\s+', "_", safe_name)
    if len(safe_name) > 30:
        hash_suffix = hashlib.md5(safe_name.encode('utf-8')).hexdigest()[:6]
        safe_name = safe_name[:23] + "_" + hash_suffix
    return safe_name

def compute_content_hash(content: str) -> str:
    return hashlib.md5(content.encode('utf-8')).hexdigest()

def backoff_delay(attempt: int, base: float = 1.0, cap: float = 30.0) -> float:
    """Exponential backoff with full jitter."""
    return random.uniform(0, min(cap, base * (2 ** attempt)))

# ---------------------------------------------------------------------------
# Circuit breaker
# ---------------------------------------------------------------------------

def _get_domain(url: str) -> str:
    return urllib.parse.urlparse(url).netloc

def check_circuit(domain: str) -> bool:
    """Returns True if circuit is open (domain should be skipped)."""
    entry = _domain_circuit.get(domain)
    if not entry:
        return False
    if entry['fail_count'] >= CIRCUIT_THRESHOLD:
        if time.monotonic() < entry['open_until']:
            return True
        # Cooldown expired — half-open: reset
        _domain_circuit[domain] = {'fail_count': 0, 'open_until': 0}
    return False

def record_domain_failure(domain: str):
    entry = _domain_circuit.setdefault(domain, {'fail_count': 0, 'open_until': 0})
    entry['fail_count'] += 1
    if entry['fail_count'] >= CIRCUIT_THRESHOLD:
        entry['open_until'] = time.monotonic() + CIRCUIT_COOLDOWN
        print(f"[circuit] {domain} circuit opened for {CIRCUIT_COOLDOWN}s after {entry['fail_count']} failures")

def record_domain_success(domain: str):
    if domain in _domain_circuit:
        _domain_circuit[domain]['fail_count'] = 0

# ---------------------------------------------------------------------------
# Rate limiting
# ---------------------------------------------------------------------------

async def acquire_rate_limit(domain: str):
    """Sleep if needed to enforce per-domain minimum interval."""
    last = _domain_last_request.get(domain, 0)
    wait = DOMAIN_MIN_INTERVAL - (time.monotonic() - last)
    if wait > 0:
        await asyncio.sleep(wait)
    _domain_last_request[domain] = time.monotonic()

# ---------------------------------------------------------------------------
# Directory setup
# ---------------------------------------------------------------------------

def setup_base_directory(start_url: str):
    """Stable output dir: scraped_data/{safe_domain}/ — no timestamp prefix."""
    parsed = urllib.parse.urlparse(start_url)
    domain = parsed.netloc.replace("www.", "")
    safe_domain = re.sub(r'[^a-zA-Z0-9_\-]', '_', domain)[:50]
    base_path = os.path.join(os.getcwd(), "scraped_data", safe_domain)
    os.makedirs(base_path, exist_ok=True)
    return base_path, domain

def setup_page_directory(base_path: str, title: str, text_only: bool = False):
    safe_title = sanitize_filename(title)
    if not safe_title:
        safe_title = "Untitled_" + datetime.now().strftime("%H%M%S")

    page_path = os.path.join(base_path, safe_title)
    counter = 1
    original_path = page_path
    while True:
        try:
            os.makedirs(page_path, exist_ok=False)
            break
        except FileExistsError:
            page_path = f"{original_path}_{counter}"
            counter += 1

    tables_path = None
    images_path = None
    if not text_only:
        tables_path = os.path.join(page_path, "tables")
        images_path = os.path.join(page_path, "images")
        os.makedirs(tables_path, exist_ok=True)
        os.makedirs(images_path, exist_ok=True)

    return page_path, tables_path, images_path

# ---------------------------------------------------------------------------
# Date filtering
# ---------------------------------------------------------------------------

_DATE_PATTERNS = [
    # 2023-05-12 or 2023/05/12 with optional time
    re.compile(r'(\d{4})[-/](\d{1,2})[-/](\d{1,2})'),
    # 2023年5月12日
    re.compile(r'(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日?'),
]

def parse_publish_date(date_str: str) -> datetime | None:
    """Parse a date string (Chinese or ISO) into an aware UTC datetime. Returns None on failure."""
    if not date_str:
        return None
    for pat in _DATE_PATTERNS:
        m = pat.search(date_str)
        if m:
            try:
                y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
                return datetime(y, mo, d, tzinfo=timezone.utc)
            except (ValueError, IndexError):
                continue
    return None

def is_within_date_window(date_str: str | None, years: int = 3) -> bool:
    """
    True if the date is within `years` of today, or if no date can be parsed
    (safe default: include unknown-date content).
    """
    if not date_str:
        return True
    dt = parse_publish_date(date_str)
    if dt is None:
        return True
    cutoff = datetime.now(timezone.utc) - timedelta(days=years * 365)
    return dt >= cutoff

# ---------------------------------------------------------------------------
# Publish info extraction
# ---------------------------------------------------------------------------

_PUBLISH_PATTERNS = [
    re.compile(r'(?:发布|创建)时间\s*[:：]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2}[日]?[\s\d:]*)'),
    re.compile(r'来源\s*[:：]\s*([^\s]+)'),
    re.compile(r'作者\s*[:：]\s*([^\s]+)'),
]

def extract_publish_info(soup) -> dict:
    """
    Extract publish date, source, and author.
    Returns a dict with keys: publish_date, source, author (all optional).
    """
    result = {}

    meta_date = soup.find('meta', attrs={'name': re.compile(r'pubdate', re.I)})
    if meta_date and meta_date.get('content'):
        result['publish_date'] = meta_date['content'].strip()

    meta_source = soup.find('meta', attrs={'name': re.compile(r'source', re.I)})
    if meta_source and meta_source.get('content'):
        result['source'] = meta_source['content'].strip()

    full_text = soup.get_text(separator=' ', strip=True)

    if 'publish_date' not in result:
        m = _PUBLISH_PATTERNS[0].search(full_text)
        if m:
            result['publish_date'] = m.group(1).strip()

    if 'source' not in result:
        m = _PUBLISH_PATTERNS[1].search(full_text)
        if m:
            result['source'] = m.group(1).strip()

    m = _PUBLISH_PATTERNS[2].search(full_text)
    if m:
        result['author'] = m.group(1).strip()

    return result

# ---------------------------------------------------------------------------
# Link discovery
# ---------------------------------------------------------------------------

def determine_static_boundary(start_url: str) -> str:
    parsed = urllib.parse.urlparse(start_url)
    ext = os.path.splitext(parsed.path)[1].lower()
    is_file = ext in ['.html', '.htm', '.php', '.jsp', '.asp', '.aspx']

    path = parsed.path
    if is_file:
        path = path.rsplit('/', 1)[0]

    path = path.rstrip('/')

    if not path:
        base_path = '/'
    else:
        parts = path.rsplit('/', 1)
        base_path = parts[0] if parts[0] else '/'
        if not base_path.endswith('/'):
            base_path += '/'

    return urllib.parse.urlunparse((parsed.scheme, parsed.netloc, base_path, '', '', ''))

def get_sub_domain_links(html: str, current_url: str, base_url: str, dynamic_root_prefix=None) -> list:
    soup = BeautifulSoup(html, 'lxml')
    links = []

    base_prefix = dynamic_root_prefix if dynamic_root_prefix else (
        base_url if base_url.endswith('/') else base_url + '/'
    )

    ignored_extensions = {
        '.zip', '.rar', '.exe', '.mp3', '.mp4', '.avi',
        '.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg'
    }

    for a_tag in soup.find_all('a', href=True):
        href = a_tag['href']
        full_url = urllib.parse.urljoin(current_url, href)
        full_url = urllib.parse.urldefrag(full_url)[0]

        parsed_url = urllib.parse.urlparse(full_url)
        ext = os.path.splitext(parsed_url.path)[1].lower()

        if parsed_url.scheme in ['http', 'https'] and ext not in ignored_extensions:
            if full_url.lower() == base_url.lower() or full_url.lower().startswith(base_prefix.lower()):
                links.append(full_url)
    return list(set(links))

# ---------------------------------------------------------------------------
# Page interaction
# ---------------------------------------------------------------------------

async def scroll_page(page):
    try:
        for _ in range(5):
            await page.evaluate("window.scrollBy(0, document.body.scrollHeight / 5)")
            await page.wait_for_timeout(1000)
        await page.evaluate("window.scrollTo(0, 0)")
        await page.wait_for_timeout(1000)
    except Exception as e:
        print(f"Warning: scroll failed: {e}")

async def intercept_route(route):
    if route.request.resource_type in ["image", "media", "font", "stylesheet"]:
        await route.abort()
    else:
        await route.continue_()

# ---------------------------------------------------------------------------
# Content extraction
# ---------------------------------------------------------------------------

def extract_main_content(html: str, fallback_soup) -> tuple:
    """
    Returns (main_html: str, title: str).
    Priority: trafilatura → readability-lxml → body fallback.
    """
    try:
        import trafilatura
        text = trafilatura.extract(html, include_tables=True, include_links=False,
                                   output_format='html', with_metadata=False)
        if text and len(BeautifulSoup(text, 'lxml').get_text(strip=True)) >= 50:
            return text, ""
    except ImportError:
        pass
    except Exception:
        pass

    try:
        doc = Document(html)
        title = doc.title() or ""
        summary = doc.summary()
        if len(BeautifulSoup(summary, 'lxml').get_text(strip=True)) >= 50:
            return summary, title
    except Exception:
        pass

    body = fallback_soup.find('body') or fallback_soup
    return str(body), ""

def clean_html_structure(soup):
    for tag in ['nav', 'footer', 'header', 'aside', 'script', 'style', 'noscript', 'iframe']:
        for match in soup.find_all(tag):
            match.decompose()

    bad = re.compile(r'menu|nav|footer|sidebar|header|banner|ad|advert|breadcrumb|share|comment', re.I)
    for tag in soup.find_all(['div', 'ul', 'ol', 'section']):
        if not hasattr(tag, 'attrs'):
            continue
        class_list = tag.attrs.get('class') if tag.attrs else None
        class_str = " ".join(class_list) if isinstance(class_list, list) else (class_list or "")
        id_str = tag.attrs.get('id', '') if tag.attrs else ''
        if bad.search(class_str) or bad.search(id_str):
            tag.decompose()

    return soup

def convert_to_markdown(html: str) -> str:
    h = html2text.HTML2Text()
    h.ignore_links = True
    h.ignore_images = True
    h.body_width = 0
    return h.handle(html)

def table_to_markdown_text(soup) -> str:
    """Convert all <table> elements in soup to inline markdown text."""
    parts = []
    for i, table in enumerate(soup.find_all('table')):
        try:
            dfs = pd.read_html(io.StringIO(str(table)))
            if dfs:
                parts.append(dfs[0].to_markdown(index=False))
        except Exception:
            parts.append(table.get_text(separator=' | ', strip=True))
        table.decompose()
    return "\n\n".join(parts)

# ---------------------------------------------------------------------------
# Table / image processing
# ---------------------------------------------------------------------------

def process_tables(soup, tables_path: str):
    tables = soup.find_all('table')
    for i, table in enumerate(tables):
        try:
            dfs = pd.read_html(io.StringIO(str(table)))
            if dfs:
                csv_path = os.path.join(tables_path, f"table_{i+1}.csv")
                dfs[0].to_csv(csv_path, index=False, encoding='utf-8-sig')
        except Exception as e:
            print(f"Error processing table {i+1}: {e}")

async def download_image(session, img_url: str, save_path: str) -> bool:
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'image/webp,image/apng,image/*,*/*;q=0.8',
        'Accept-Encoding': 'gzip, deflate, br',
        'Referer': img_url,
    }
    for attempt in range(3):
        try:
            async with session.get(img_url, headers=headers, timeout=15, ssl=False) as response:
                if response.status == 200:
                    content = await response.read()
                    with open(save_path, 'wb') as f:
                        f.write(content)
                    return True
                elif response.status in [301, 302, 303, 307, 308]:
                    redirect_url = response.headers.get('Location')
                    if redirect_url:
                        if not redirect_url.startswith('http'):
                            redirect_url = urllib.parse.urljoin(img_url, redirect_url)
                        img_url = redirect_url
                        continue
                else:
                    print(f"Failed to download {img_url}: HTTP {response.status}")
                    break
        except Exception as e:
            if attempt == 2:
                print(f"Error downloading {img_url} after 3 attempts: {e}")
            await asyncio.sleep(backoff_delay(attempt))
    return False

async def process_images(soup, base_url: str, images_path: str, session):
    download_tasks = []
    image_count = 0

    for img in soup.find_all('img'):
        src = None
        for attr in ['data-src', 'data-original', 'data-url', 'data-echo', 'data-lazy-src', 'src']:
            val = img.get(attr)
            if val and isinstance(val, str) and val.strip():
                if val.strip().startswith('data:image'):
                    continue
                src = val.strip()
                break

        if src:
            img_url = urllib.parse.urljoin(base_url, src)
            if img_url.startswith('data:'):
                continue

            parsed_img_url = urllib.parse.urlparse(img_url)
            ext = os.path.splitext(parsed_img_url.path)[1]
            if not ext or ext.lower() not in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg']:
                ext = '.jpg'

            image_count += 1
            filename = f"image_{image_count}{ext}"
            save_path = os.path.join(images_path, filename)

            img['src'] = f"./images/{filename}"
            for attr in ['data-src', 'data-original', 'data-url', 'data-echo', 'data-lazy-src']:
                if img.get(attr):
                    del img[attr]

            download_tasks.append(download_image(session, img_url, save_path))

    for video in soup.find_all('video'):
        poster = video.get('poster')
        if poster and isinstance(poster, str) and poster.strip():
            img_url = urllib.parse.urljoin(base_url, poster.strip())
            if not img_url.startswith('data:'):
                parsed_img_url = urllib.parse.urlparse(img_url)
                ext = os.path.splitext(parsed_img_url.path)[1]
                if not ext or ext.lower() not in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
                    ext = '.jpg'
                image_count += 1
                filename = f"image_{image_count}{ext}"
                save_path = os.path.join(images_path, filename)
                video['poster'] = f"./images/{filename}"
                download_tasks.append(download_image(session, img_url, save_path))

    url_pattern = re.compile(r'url\(\s*[\'\"]?(.*?)[\'\"]?\s*\)')
    for tag in soup.find_all(style=True):
        style_content = tag['style']
        if 'background' in style_content:
            matches = url_pattern.findall(style_content)
            new_style = style_content
            for match in matches:
                if match.startswith('data:'):
                    continue
                img_url = urllib.parse.urljoin(base_url, match)
                parsed_img_url = urllib.parse.urlparse(img_url)
                ext = os.path.splitext(parsed_img_url.path)[1]
                if not ext or ext.lower() not in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
                    ext = '.jpg'
                image_count += 1
                filename = f"image_{image_count}{ext}"
                save_path = os.path.join(images_path, filename)
                new_style = new_style.replace(match, f"./images/{filename}")
                download_tasks.append(download_image(session, img_url, save_path))
                if tag.name in ['xg-poster', 'div', 'span'] and 'poster' in tag.get('class', []):
                    new_img = soup.new_tag('img', src=f"./images/{filename}")
                    tag.append(new_img)
            tag['style'] = new_style

    if download_tasks:
        await asyncio.gather(*download_tasks)

# ---------------------------------------------------------------------------
# File download helpers
# ---------------------------------------------------------------------------

def get_cookies_for_url(pw_cookies, url: str) -> dict | None:
    if not pw_cookies:
        return None
    parsed_url = urllib.parse.urlparse(url)
    domain = parsed_url.hostname
    if not domain:
        return None
    valid_cookies = {}
    for cookie in pw_cookies:
        c_domain = cookie['domain']
        if (c_domain == domain or
                (c_domain.startswith('.') and domain.endswith(c_domain[1:])) or
                domain == c_domain.lstrip('.')):
            valid_cookies[cookie['name']] = cookie['value']
    return valid_cookies if valid_cookies else None

async def download_file(url: str, save_path: str, user_agent: str, cookies=None) -> bool:
    headers = {
        'User-Agent': user_agent,
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
        'Upgrade-Insecure-Requests': '1',
    }
    for attempt in range(3):
        try:
            connector = aiohttp.TCPConnector(ssl=False)
            async with aiohttp.ClientSession(connector=connector, headers=headers, cookies=cookies) as session:
                async with session.get(url, timeout=300) as response:
                    if response.status == 200:
                        with open(save_path, 'wb') as f:
                            async for chunk in response.content.iter_chunked(8192):
                                f.write(chunk)
                        return True
                    else:
                        print(f"Download {url} failed: HTTP {response.status} (attempt {attempt+1})")
        except Exception as e:
            print(f"Error downloading {url} (attempt {attempt+1}): {e}")
            await asyncio.sleep(backoff_delay(attempt))
    return False

async def extract_text_from_pdf(file_path: str) -> str | None:
    """Extract text from a PDF file using pdfplumber (runs in thread)."""
    def _extract():
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n\n".join(pages) if pages else None
        except ImportError:
            print("pdfplumber not installed — PDF text extraction skipped")
            return None
        except Exception as e:
            print(f"PDF extraction error: {e}")
            return None

    return await asyncio.to_thread(_extract)

async def extract_text_from_docx(file_path: str) -> str | None:
    """Extract text from .docx (or attempt LibreOffice conversion for .doc)."""
    def _extract():
        path = file_path
        converted = False

        if path.lower().endswith('.doc') and not path.lower().endswith('.docx'):
            try:
                import subprocess, tempfile, shutil
                tmpdir = tempfile.mkdtemp()
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', tmpdir, path],
                    capture_output=True, timeout=60
                )
                if result.returncode == 0:
                    base = os.path.splitext(os.path.basename(path))[0]
                    converted_path = os.path.join(tmpdir, base + '.docx')
                    if os.path.exists(converted_path):
                        path = converted_path
                        converted = True
            except Exception as e:
                print(f".doc conversion warning: {e}")

        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            if converted:
                import shutil
                shutil.rmtree(os.path.dirname(path), ignore_errors=True)
            return text if text else None
        except ImportError:
            print("python-docx not installed — Word text extraction skipped")
            return None
        except Exception as e:
            print(f"Docx extraction error: {e}")
            return None

    return await asyncio.to_thread(_extract)

def save_as_markdown_from_text(text: str, title: str, source_url: str,
                                page_path: str, publish_info: dict | None = None) -> str:
    """Write extracted document text as content.md with YAML frontmatter."""
    header = "---\n"
    header += f'title: "{title}"\n'
    header += f'source_url: "{source_url}"\n'
    if publish_info:
        if publish_info.get('publish_date'):
            header += f'publish_date: "{publish_info["publish_date"]}"\n'
        if publish_info.get('source'):
            header += f'source: "{publish_info["source"]}"\n'
        if publish_info.get('author'):
            header += f'author: "{publish_info["author"]}"\n'
    header += "---\n\n"
    header += f"# {title}\n\n"

    content = header + text
    md_path = os.path.join(page_path, "content.md")
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(content)
    return content

# ---------------------------------------------------------------------------
# Pause / stop
# ---------------------------------------------------------------------------

async def check_pause_stop(task_id: str) -> bool:
    events = task_events.get(task_id)
    if not events:
        return False

    if events['stop'].is_set():
        await asyncio.to_thread(db_manager.update_task_status, task_id, 'stopped')
        return True

    if not events['pause'].is_set():
        await asyncio.to_thread(db_manager.update_task_status, task_id, 'paused')
        await events['pause'].wait()

        if events['stop'].is_set():
            await asyncio.to_thread(db_manager.update_task_status, task_id, 'stopped')
            return True

        await asyncio.to_thread(db_manager.update_task_status, task_id, 'running')

    return False

# ---------------------------------------------------------------------------
# Browser config
# ---------------------------------------------------------------------------

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36 Edg/121.0.0.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:123.0) Gecko/20100101 Firefox/123.0",
]

VIEWPORTS = [
    {"width": 1920, "height": 1080},
    {"width": 1440, "height": 900},
    {"width": 1536, "height": 864},
    {"width": 1366, "height": 768},
]

# ---------------------------------------------------------------------------
# Core URL processor
# ---------------------------------------------------------------------------

async def process_single_url(task_id: str, current_url: str, start_url: str,
                              base_path: str, browser,
                              text_only: bool = False,
                              date_filter: bool = False,
                              update_mode: bool = False) -> None:
    """Process one URL: download, extract text, save markdown."""

    if await check_pause_stop(task_id):
        return

    domain = _get_domain(current_url)

    # Circuit breaker check
    if check_circuit(domain):
        print(f"[circuit] Skipping {current_url} — {domain} circuit open")
        await asyncio.to_thread(db_manager.mark_url_filtered, task_id, current_url, "circuit_open")
        return

    # Rate limiting
    await acquire_rate_limit(domain)

    # -----------------------------------------------------------------------
    # Document file branch (PDF / Word / Excel / PPT)
    # -----------------------------------------------------------------------
    parsed_url = urllib.parse.urlparse(current_url)
    ext = os.path.splitext(parsed_url.path)[1].lower()

    if ext in ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx']:
        filename = os.path.basename(parsed_url.path) or f"download_{random.randint(1000,9999)}{ext}"
        safe_title = re.sub(r'[\\/*?:"<>|]', "", filename)

        page_path, _, _ = await asyncio.to_thread(setup_page_directory, base_path, safe_title, True)
        save_path = os.path.join(page_path, filename)

        ua = random.choice(USER_AGENTS)
        cookie_dict = None
        try:
            if browser.contexts:
                pw_cookies = await browser.contexts[0].cookies()
                cookie_dict = get_cookies_for_url(pw_cookies, current_url)
        except Exception:
            pass

        success = await download_file(current_url, save_path, ua, cookies=cookie_dict)
        if not success:
            record_domain_failure(domain)
            await asyncio.to_thread(db_manager.mark_url_failed, task_id, current_url, "File download failed")
            return

        # Convert text-extractable formats to markdown
        extracted_text = None
        if ext == '.pdf':
            extracted_text = await extract_text_from_pdf(save_path)
        elif ext in ['.doc', '.docx']:
            extracted_text = await extract_text_from_docx(save_path)

        if extracted_text:
            content = save_as_markdown_from_text(extracted_text, filename, current_url, page_path)
            # Remove binary file after successful text extraction
            try:
                os.remove(save_path)
            except Exception:
                pass
            content_hash = compute_content_hash(content)
            record_domain_success(domain)
            await asyncio.to_thread(
                db_manager.mark_url_scraped, task_id, current_url, filename, page_path, 'article',
                content_hash
            )
        else:
            # Keep binary for non-text-extractable formats
            record_domain_success(domain)
            await asyncio.to_thread(
                db_manager.mark_url_scraped, task_id, current_url, filename, page_path, 'article'
            )
        return

    # -----------------------------------------------------------------------
    # Normal HTML processing
    # -----------------------------------------------------------------------
    context = None
    page = None
    try:
        ua = random.choice(USER_AGENTS)
        vp = random.choice(VIEWPORTS)

        if browser.contexts:
            context = browser.contexts[0]
        else:
            context = await browser.new_context(
                user_agent=ua,
                viewport=vp,
                ignore_https_errors=True
            )

        page = await context.new_page()
        await Stealth().apply_stealth_async(page)
        await page.route("**/*", intercept_route)

        for attempt in range(3):
            try:
                await page.goto(current_url, wait_until="domcontentloaded", timeout=60000)
                await page.wait_for_selector('body', state='attached', timeout=30000)
                if await check_pause_stop(task_id):
                    return
                await page.wait_for_timeout(3000)
                break
            except asyncio.CancelledError:
                raise
            except Exception as goto_e:
                print(f"Warning: goto error for {current_url} (attempt {attempt+1}): {goto_e}")
                if attempt == 2:
                    record_domain_failure(domain)
                    print(f"Proceeding with loaded content after {attempt+1} failures.")
                else:
                    await asyncio.sleep(backoff_delay(attempt))

        await scroll_page(page)
        html_content = await page.content()

        # Determine crawl boundary
        dynamic_root = task_root_prefixes.get(task_id)
        if not dynamic_root:
            task_data = await asyncio.to_thread(db_manager.get_task, task_id)
            if task_data and task_data.get('dynamic_root'):
                dynamic_root = task_data['dynamic_root']
                task_root_prefixes[task_id] = dynamic_root

        if current_url == start_url and not dynamic_root:
            dynamic_root = determine_static_boundary(start_url)
            task_root_prefixes[task_id] = dynamic_root
            await asyncio.to_thread(db_manager.update_task_dynamic_root, task_id, dynamic_root)

        new_links = get_sub_domain_links(html_content, current_url, start_url, dynamic_root)
        if new_links:
            await asyncio.to_thread(db_manager.add_discovered_urls, task_id, current_url, new_links)

        full_soup = BeautifulSoup(html_content, 'lxml')

        # Extract title
        title = full_soup.title.string if full_soup.title and full_soup.title.string else "Untitled Page"

        # === DATE FILTER (must happen before creating directories) ===
        publish_info = extract_publish_info(full_soup)

        if date_filter and not is_within_date_window(publish_info.get('publish_date'), years=3):
            print(f"[date-filter] Skipping {current_url} — date: {publish_info.get('publish_date')}")
            await asyncio.to_thread(
                db_manager.mark_url_filtered, task_id, current_url,
                f"date_out_of_window:{publish_info.get('publish_date')}"
            )
            return

        # Create page directory
        page_path, tables_path, images_path = await asyncio.to_thread(
            setup_page_directory, base_path, title, text_only
        )

        body_soup = full_soup.find('body') or full_soup

        # Process images only when not in text-only mode
        if not text_only:
            pw_cookies = await context.cookies()
            connector = aiohttp.TCPConnector(ssl=False)
            headers = {'User-Agent': ua}
            async with aiohttp.ClientSession(connector=connector, headers=headers) as session:
                session._pw_cookies = pw_cookies
                await process_images(body_soup, current_url, images_path, session)

        # Extract main content via trafilatura → readability → body fallback
        cleaned_soup = clean_html_structure(BeautifulSoup(str(full_soup), 'lxml'))
        cleaned_html = str(cleaned_soup)
        main_html, extracted_title = extract_main_content(cleaned_html, cleaned_soup)
        if extracted_title:
            title = extracted_title
        main_soup = BeautifulSoup(main_html, 'lxml')

        # Tables: CSV in normal mode, inline markdown in text-only mode
        table_extra = ""
        if text_only:
            table_extra = table_to_markdown_text(main_soup)
        elif tables_path:
            process_tables(main_soup, tables_path)

        # Generate markdown
        markdown_content = convert_to_markdown(str(main_soup))
        if table_extra:
            markdown_content = table_extra + "\n\n" + markdown_content

        # Build YAML frontmatter
        header = "---\n"
        header += f'title: "{title}"\n'
        header += f'source_url: "{current_url}"\n'
        if publish_info.get('publish_date'):
            header += f'publish_date: "{publish_info["publish_date"]}"\n'
        if publish_info.get('source'):
            header += f'source: "{publish_info["source"]}"\n'
        if publish_info.get('author'):
            header += f'author: "{publish_info["author"]}"\n'
        header += "---\n\n"
        header += f"# {title}\n\n"

        markdown_content = header + markdown_content

        # Content type heuristic
        content_type = 'node'
        for block in full_soup.find_all(['p', 'div', 'span', 'article', 'section']):
            if not hasattr(block, 'attrs'):
                continue
            temp_block = BeautifulSoup(str(block), 'lxml')
            for a in temp_block.find_all('a'):
                a.decompose()
            if len(temp_block.get_text(strip=True)) > 30:
                content_type = 'article'
                break

        new_hash = compute_content_hash(markdown_content)

        # Iterative update: skip file write if content unchanged
        if update_mode:
            old_hash = await asyncio.to_thread(db_manager.get_url_content_hash, task_id, current_url)
            if old_hash and old_hash == new_hash:
                print(f"[update] Unchanged: {current_url}")
                await asyncio.to_thread(
                    db_manager.mark_url_scraped, task_id, current_url, title, page_path, content_type, new_hash
                )
                record_domain_success(domain)
                return

        md_path = os.path.join(page_path, "content.md")
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)

        record_domain_success(domain)
        await asyncio.to_thread(
            db_manager.mark_url_scraped, task_id, current_url, title, page_path, content_type, new_hash
        )

    except Exception as page_e:
        import traceback
        traceback.print_exc()
        print(f"Error scraping {current_url}: {page_e}")
        record_domain_failure(domain)
        await asyncio.to_thread(db_manager.mark_url_failed, task_id, current_url, str(page_e))
    finally:
        if page:
            await page.close()
        if context and not browser.contexts:
            await context.close()

# ---------------------------------------------------------------------------
# Worker task
# ---------------------------------------------------------------------------

async def worker_task(task_id: str, start_url: str, base_path: str, browser,
                      semaphore: asyncio.Semaphore,
                      text_only: bool = False,
                      date_filter: bool = False,
                      update_mode: bool = False):
    async with semaphore:
        try:
            should_stop = await check_pause_stop(task_id)
            if should_stop:
                return False

            current_url = await asyncio.to_thread(db_manager.get_and_lock_pending_url, task_id)
            if not current_url:
                return True

            delay = random.uniform(1.0, 3.0)
            await asyncio.sleep(delay)

            await process_single_url(
                task_id, current_url, start_url, base_path, browser,
                text_only=text_only, date_filter=date_filter, update_mode=update_mode
            )
            return True
        except asyncio.CancelledError:
            print(f"Worker task cancelled for task {task_id}.")
            raise
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"Worker task crashed: {e}")
            return True

# ---------------------------------------------------------------------------
# Crawl worker (main background task)
# ---------------------------------------------------------------------------

async def crawl_worker(task_id: str, start_url: str, headless: bool = True,
                       text_only: bool = False, date_filter: bool = False,
                       update_mode: bool = False):
    """Background manager that schedules concurrent workers."""

    init_task_events(task_id)
    await asyncio.to_thread(db_manager.reset_processing_urls, task_id)

    base_path, base_domain = setup_base_directory(start_url)

    # Persist stable base_path to DB
    await asyncio.to_thread(db_manager.create_task, task_id, start_url, start_url)
    await asyncio.to_thread(db_manager.update_task_base_path, task_id, base_path)

    semaphore = asyncio.Semaphore(5)
    os.environ["PLAYWRIGHT_BROWSERS_PATH"] = "0"

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(
                headless=headless,
                args=["--disable-blink-features=AutomationControlled"]
            )

            print(f"[{task_id}] Session warm-up on {start_url}...")
            shared_context = await browser.new_context(
                user_agent=random.choice(USER_AGENTS),
                viewport=random.choice(VIEWPORTS),
                ignore_https_errors=True
            )
            try:
                warmup_page = await shared_context.new_page()
                await Stealth().apply_stealth_async(warmup_page)
                await warmup_page.route("**/*", intercept_route)
                await warmup_page.goto(start_url, wait_until="domcontentloaded", timeout=45000)
                await scroll_page(warmup_page)
                await warmup_page.wait_for_timeout(2000)
                print(f"[{task_id}] Warm-up completed.")
                await warmup_page.close()
            except Exception as e:
                print(f"[{task_id}] Warm-up error (continuing): {e}")

            active_tasks = set()
            while True:
                should_stop = await check_pause_stop(task_id)
                if should_stop:
                    break

                done = {t for t in active_tasks if t.done()}
                active_tasks.difference_update(done)

                stop_signal = False
                for t in done:
                    if t.cancelled():
                        continue
                    if t.exception() is not None:
                        print(f"Task exception: {t.exception()}")
                        continue
                    if t.result() is False:
                        stop_signal = True
                        break

                if stop_signal:
                    break

                active_count = await asyncio.to_thread(db_manager.get_active_count, task_id)
                if active_count == 0 and len(active_tasks) == 0:
                    await asyncio.to_thread(db_manager.update_task_status, task_id, 'completed')
                    break

                while len(active_tasks) < 5 and active_count > 0:
                    task = asyncio.create_task(
                        worker_task(
                            task_id, start_url, base_path, browser, semaphore,
                            text_only=text_only, date_filter=date_filter, update_mode=update_mode
                        )
                    )
                    active_tasks.add(task)
                    active_count -= 1

                if active_tasks:
                    stop_event_task = asyncio.create_task(task_events[task_id]['stop'].wait())
                    wait_tasks = list(active_tasks) + [stop_event_task]

                    done, _ = await asyncio.wait(wait_tasks, return_when=asyncio.FIRST_COMPLETED)

                    if stop_event_task in done:
                        await asyncio.to_thread(db_manager.update_task_status, task_id, 'stopped')
                        break
                    else:
                        stop_event_task.cancel()
                else:
                    await asyncio.sleep(1)

    except Exception as e:
        print(f"Crawl fatally failed: {e}")
        await asyncio.to_thread(db_manager.update_task_status, task_id, 'failed')
    finally:
        if 'active_tasks' in locals() and active_tasks:
            for t in active_tasks:
                if not t.done():
                    t.cancel()
            await asyncio.wait(active_tasks, timeout=2.0)

        await asyncio.to_thread(db_manager.reset_processing_urls, task_id)
        if task_id in task_events:
            del task_events[task_id]
        if task_id in task_root_prefixes:
            del task_root_prefixes[task_id]
